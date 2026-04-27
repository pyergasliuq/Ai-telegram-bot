from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from pathlib import Path

log = logging.getLogger(__name__)


SUPPORTED_INPUT_EXT: tuple[str, ...] = (".pdf", ".docx", ".txt", ".md", ".markdown", ".rtf")
SUPPORTED_OUTPUT_FMT: tuple[str, ...] = ("md", "txt", "docx", "pdf")
MAX_INPUT_BYTES: int = 8 * 1024 * 1024
MAX_INPUT_CHARS: int = 200_000


@dataclass
class IngestedFile:
    name: str
    text: str
    chars: int


class FileError(Exception):
    pass


def ingest(name: str, data: bytes) -> IngestedFile:
    if len(data) > MAX_INPUT_BYTES:
        raise FileError("file_too_large")
    suffix = Path(name).suffix.lower()
    if suffix not in SUPPORTED_INPUT_EXT:
        raise FileError("unsupported_format")
    if suffix == ".pdf":
        text = _read_pdf(data)
    elif suffix == ".docx":
        text = _read_docx(data)
    else:
        text = data.decode("utf-8", errors="ignore")
    text = text.strip()
    if not text:
        raise FileError("empty_file")
    if len(text) > MAX_INPUT_CHARS:
        text = text[:MAX_INPUT_CHARS]
    return IngestedFile(name=name, text=text, chars=len(text))


def _read_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as e:
        raise FileError(f"pdf_lib_unavailable:{e}") from e
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(p for p in parts if p)
    except Exception as e:
        raise FileError(f"pdf_parse_error:{e}") from e


def _read_docx(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as e:
        raise FileError(f"docx_lib_unavailable:{e}") from e
    try:
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        raise FileError(f"docx_parse_error:{e}") from e


def render(text: str, fmt: str, *, title: str = "") -> tuple[bytes, str]:
    fmt = fmt.lower().strip()
    if fmt in ("md", "markdown"):
        return text.encode("utf-8"), "md"
    if fmt == "txt":
        return text.encode("utf-8"), "txt"
    if fmt == "docx":
        return _render_docx(text, title=title), "docx"
    if fmt == "pdf":
        return _render_pdf(text, title=title), "pdf"
    raise FileError(f"unsupported_output:{fmt}")


def _render_docx(text: str, *, title: str = "") -> bytes:
    try:
        from docx import Document
    except Exception as e:
        raise FileError(f"docx_lib_unavailable:{e}") from e
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        doc.add_paragraph(block)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_pdf(text: str, *, title: str = "") -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as e:
        raise FileError(f"pdf_lib_unavailable:{e}") from e
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    flow: list = []
    if title:
        flow.append(Paragraph(title, styles["Title"]))
        flow.append(Spacer(1, 12))
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        safe = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        flow.append(Paragraph(safe.replace("\n", "<br/>"), styles["BodyText"]))
        flow.append(Spacer(1, 8))
    doc.build(flow)
    return buf.getvalue()
